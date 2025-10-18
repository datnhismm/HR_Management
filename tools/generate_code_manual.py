"""Generate a developer manual (.docx) for the hr_management_app package.

This tool will:
- Walk the `hr_management_app` package and parse all .py files.
- Use the AST to extract top-level classes and functions, their signatures, and docstrings.
- Generate templated explanations for each function/class (purpose, inputs, outputs, side-effects, exceptions).
- Search for known deprecation hotspots (utcnow usage) and include them.
- Run pytest and attach a test summary section.
- Save the manual to reports/system_manual_<ts>.docx

Usage: python tools/generate_code_manual.py

Note: This generator uses simple heuristics to create the explanations. For the most accurate
manual, review and adjust the generated document.
"""
import ast
import os
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    print("Missing dependency: python-docx. Install with 'pip install python-docx'")
    raise


ROOT = Path(__file__).resolve().parents[1]
PKG_ROOT = ROOT / 'hr_management_app'


def find_py_files():
    files = [p for p in PKG_ROOT.rglob('*.py') if p.is_file()]
    return sorted(files)


def parse_file(filepath: Path):
    src = filepath.read_text(encoding='utf-8')
    tree = ast.parse(src)
    return tree, src


def extract_definitions(tree):
    defs = []
    for node in tree.body:
        if isinstance(node, ast.FunctionDef):
            defs.append(('function', node))
        elif isinstance(node, ast.ClassDef):
            defs.append(('class', node))
    return defs


def signature_from_function(node: ast.FunctionDef):
    # Build a simple signature string
    args = []
    # positional args
    for a in node.args.args:
        if a.arg == 'self':
            args.append('self')
            continue
        args.append(a.arg)
    # vararg
    if node.args.vararg:
        args.append('*' + node.args.vararg.arg)
    # keyword-only args
    for a in node.args.kwonlyargs:
        args.append(a.arg)
    # kwarg
    if node.args.kwarg:
        args.append('**' + node.args.kwarg.arg)
    sig = f"def {node.name}({', '.join(args)})"
    return sig


def extract_params(node: ast.FunctionDef):
    params = []
    # positional and keyword-only args
    all_args = list(node.args.args) + list(node.args.kwonlyargs)
    defaults = [None] * (len(node.args.args) - len(node.args.defaults)) + node.args.defaults
    # map defaults to positional args
    pos_defaults = defaults[: len(node.args.args)] if node.args.defaults else [None] * len(node.args.args)
    for i, a in enumerate(node.args.args):
        if a.arg == 'self':
            continue
        default = None
        if node.args.defaults and i >= len(node.args.args) - len(node.args.defaults):
            # compute default index
            didx = i - (len(node.args.args) - len(node.args.defaults))
            try:
                default_node = node.args.defaults[didx]
                default = ast.unparse(default_node) if hasattr(ast, 'unparse') else None
            except Exception:
                default = None
        params.append({'name': a.arg, 'default': default})
    # kwonly defaults
    for i, a in enumerate(node.args.kwonlyargs):
        default = None
        try:
            default_node = node.args.kw_defaults[i]
            if default_node is not None:
                default = ast.unparse(default_node) if hasattr(ast, 'unparse') else None
        except Exception:
            default = None
        params.append({'name': a.arg, 'default': default})
    # vararg/kwarg
    if node.args.vararg:
        params.append({'name': '*' + node.args.vararg.arg, 'default': None})
    if node.args.kwarg:
        params.append({'name': '**' + node.args.kwarg.arg, 'default': None})
    return params


def find_returns_and_raises(node: ast.FunctionDef):
    returns = False
    raises = []
    for n in ast.walk(node):
        if isinstance(n, ast.Return):
            returns = True
        if isinstance(n, ast.Raise):
            try:
                if n.exc:
                    exc = ast.unparse(n.exc) if hasattr(ast, 'unparse') else None
                    raises.append(exc)
            except Exception:
                raises.append(None)
    return returns, raises


def signature_from_class(node: ast.ClassDef):
    # find base names
    bases = [ast.unparse(b) if hasattr(ast, 'unparse') else getattr(b, 'id', '?') for b in node.bases]
    if bases:
        return f"class {node.name}({', '.join(bases)})"
    return f"class {node.name}"


def summarize_docstring(doc: str) -> str:
    if not doc:
        return ''
    # take first paragraph
    parts = doc.strip().split('\n\n')
    return parts[0].strip()


def find_utcnow_hotspots(src: str):
    hits = []
    for m in re.finditer(r"utcnow\(|datetime\.utcnow\(", src):
        ln = src[:m.start()].count('\n') + 1
        line = src.splitlines()[ln - 1].strip()
        hits.append((ln, line))
    return hits


def run_pytest():
    print('Running pytest...')
    proc = subprocess.run([sys.executable, '-m', 'pytest', '-q'], cwd=ROOT, capture_output=True, text=True)
    return proc.returncode, proc.stdout + '\n' + proc.stderr


def build_manual(outpath: Path, files_info, pytest_output):
    doc = Document()
    doc.add_heading('HR_Management Developer Manual', level=1)
    doc.add_paragraph(f'Generated: {datetime.now(timezone.utc).astimezone().isoformat()}')

    doc.add_heading('Table of Contents', level=2)
    for f in files_info:
        doc.add_paragraph(str(f['relpath']), style='List Number')

    # File map (quick index)
    doc.add_heading('File Map', level=2)
    for f in files_info:
        doc.add_paragraph(f"{f['relpath']} - {len(f['items'])} top-level defs")

    for f in files_info:
        doc.add_page_break()
        doc.add_heading(f['relpath'], level=2)
        doc.add_paragraph(f['summary'] or 'No file-level docstring available.')

        if f['hotspots']:
            doc.add_paragraph('Deprecation / Hotspots:', style='IntenseQuote')
            for ln, line in f['hotspots']:
                doc.add_paragraph(f'Line {ln}: {line}', style='Quote')

        for item in f['items']:
            if item['type'] == 'function':
                doc.add_heading(item['name'], level=3)
                doc.add_paragraph(item['signature'])
                # parameters
                try:
                    params = item.get('params') or []
                    if params:
                        doc.add_paragraph('Parameters:')
                        for p in params:
                            doc.add_paragraph(f"- {p['name']} (default={p['default']})")
                except Exception:
                    pass
                # returns / raises
                try:
                    returns = item.get('returns')
                    raises = item.get('raises') or []
                    doc.add_paragraph(f"Returns: {'Yes' if returns else 'None/implicit'}")
                    if raises:
                        doc.add_paragraph('Raises:')
                        for r in raises:
                            doc.add_paragraph(f"- {r}")
                except Exception:
                    pass
                if item['doc']:
                    doc.add_paragraph('Docstring:')
                    doc.add_paragraph(item['doc'])
                doc.add_paragraph('Summary:')
                doc.add_paragraph(item['summary'] or 'No summary available.')
                # heuristics for IO/side-effects
                side = []
                src = item.get('source', '')
                if 'send_email' in src or 'smtplib' in src:
                    side.append('Sends email (side-effect)')
                if 'open(' in src or 'os.startfile' in src or 'shutil' in src:
                    side.append('Performs file I/O or opens files')
                if 'INSERT INTO' in src or 'UPDATE ' in src or 'DELETE FROM' in src:
                    side.append('Modifies database')
                if side:
                    doc.add_paragraph('Side-effects: ' + '; '.join(side))
                # example usage template
                doc.add_paragraph('Example usage:')
                ex = f"{item['signature']}\n# TODO: add example usage"
                doc.add_paragraph(ex)

            elif item['type'] == 'class':
                doc.add_heading(item['name'], level=3)
                doc.add_paragraph(item['signature'])
                if item['doc']:
                    doc.add_paragraph('Docstring:')
                    doc.add_paragraph(item['doc'])
                doc.add_paragraph('Summary:')
                doc.add_paragraph(item['summary'] or 'No summary available.')

    # append pytest output
    doc.add_page_break()
    doc.add_heading('Test Summary', level=2)
    if pytest_output:
        snippet = pytest_output.strip()
        if len(snippet) > 4000:
            snippet = snippet[:4000] + '\n... (truncated)'
        p = doc.add_paragraph(snippet)
        try:
            # some python-docx installations may return None for p.style
            if p.style is not None:
                p.style.font.size = Pt(9)
        except Exception:
            pass

    outpath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(outpath))


def main():
    files = find_py_files()
    files_info = []
    for p in files:
        rel = p.relative_to(ROOT)
        tree, src = parse_file(p)
        docstring = ast.get_docstring(tree)
        items = []
        for typ, node in extract_definitions(tree):
            if typ == 'function':
                sig = signature_from_function(node)
                doc = ast.get_docstring(node) or ''
                summary = summarize_docstring(doc)
                src_segment = '\n'.join(src.splitlines()[node.lineno - 1: node.end_lineno]) if hasattr(node, 'end_lineno') else ''
                params = extract_params(node)
                returns, raises = find_returns_and_raises(node)
                items.append({'type': 'function', 'name': node.name, 'signature': sig, 'doc': doc, 'summary': summary, 'source': src_segment, 'params': params, 'returns': returns, 'raises': raises})
            elif typ == 'class':
                sig = signature_from_class(node)
                doc = ast.get_docstring(node) or ''
                summary = summarize_docstring(doc)
                items.append({'type': 'class', 'name': node.name, 'signature': sig, 'doc': doc, 'summary': summary})
        hotspots = find_utcnow_hotspots(src)
        files_info.append({'relpath': str(rel), 'summary': docstring, 'items': items, 'hotspots': hotspots})

    rc, pytest_out = run_pytest()

    ts = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')
    outpath = ROOT / 'reports' / f'system_manual_{ts}.docx'
    print('Building manual to', outpath)
    build_manual(outpath, files_info, pytest_out)
    print('Manual written to', outpath)


if __name__ == '__main__':
    main()
