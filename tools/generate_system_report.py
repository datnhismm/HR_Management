"""Generate a system report and save it as a .docx file.

This script will:
- run pytest (quiet) and capture the output
- search the codebase for known deprecation hotspots (datetime.utcnow usages)
- produce a Word (.docx) file under reports/ with a timestamped name

Requires: python-docx (already listed in requirements.txt)
Usage: python tools/generate_system_report.py
"""
import os
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    print("Missing dependency: python-docx. Install with 'pip install python-docx'")
    raise


ROOT = Path(__file__).resolve().parents[1]


def run_pytest():
    print("Running pytest...")
    # run pytest and capture output
    proc = subprocess.run([sys.executable, "-m", "pytest", "-q"], cwd=ROOT, capture_output=True, text=True)
    return proc.returncode, proc.stdout + "\n" + proc.stderr


def find_deprecation_hotspots():
    pattern = re.compile(r"utcnow\(|datetime\.utcnow\(")
    hits = []
    for p in ROOT.rglob("*.py"):
        try:
            txt = p.read_text(encoding="utf-8")
        except Exception:
            continue
        for m in pattern.finditer(txt):
            # find line number
            line_no = txt[: m.start()].count("\n") + 1
            hits.append((str(p.relative_to(ROOT)), line_no, txt.splitlines()[line_no - 1].strip()))
    return hits


def build_docx(report_path: Path, pytest_output: str, hotspots):
    doc = Document()
    doc.add_heading('HR_Management System Report', level=1)
    ts = datetime.now(timezone.utc).astimezone().isoformat()
    doc.add_paragraph(f'Generated: {ts}')

    doc.add_heading('Test Summary', level=2)
    # include first 2000 chars of pytest output then attach a code block for full output
    para = doc.add_paragraph()
    para.add_run('Pytest output (truncated):').bold = True
    out_snip = pytest_output.strip()
    if len(out_snip) > 2000:
        out_short = out_snip[:2000] + '\n... (truncated)'
    else:
        out_short = out_snip
    p = doc.add_paragraph(out_short)
    p.style.font.size = Pt(9)

    doc.add_heading('Deprecation / Hotspots', level=2)
    if not hotspots:
        doc.add_paragraph('No utcnow() hotspots found.')
    else:
        tbl = doc.add_table(rows=1, cols=3)
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = 'File'
        hdr_cells[1].text = 'Line'
        hdr_cells[2].text = 'Code'
        for f, ln, code in hotspots:
            row = tbl.add_row().cells
            row[0].text = f
            row[1].text = str(ln)
            row[2].text = code

    doc.add_heading('Notes & Next Steps', level=2)
    doc.add_paragraph('- Replace datetime.utcnow() with datetime.now(timezone.utc) where appropriate.')
    doc.add_paragraph('- Consider adding CI job to run GUI tests under Windows or xvfb (Linux).')
    doc.add_paragraph('- Add mocked-SMTP unit tests for outbox processing.')

    doc.save(report_path)


def main():
    rc, output = run_pytest()
    hotspots = find_deprecation_hotspots()
    reports_dir = ROOT / 'reports'
    reports_dir.mkdir(exist_ok=True)
    ts = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')
    report_path = reports_dir / f'system_report_{ts}.docx'
    print(f'Writing report to {report_path}')
    build_docx(report_path, output, hotspots)
    print('Done. Report saved at:', report_path)
    if rc != 0:
        print('pytest returned non-zero exit code:', rc)


if __name__ == '__main__':
    main()
