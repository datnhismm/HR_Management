import os
from typing import List, Dict, Any
import pandas as pd

def parse_excel(path: str) -> List[Dict[str, Any]]:
    """Read an Excel file and return list of row dicts (columns as-is)."""
    xls = pd.read_excel(path, dtype=str)
    records = []
    for _, row in xls.fillna("").iterrows():
        records.append({str(k).strip(): (v.strip() if isinstance(v, str) else v) for k, v in row.items()})
    return records

def parse_csv(path: str) -> List[Dict[str, Any]]:
    df = pd.read_csv(path, dtype=str)
    records = []
    for _, row in df.fillna("").iterrows():
        records.append({str(k).strip(): (v.strip() if isinstance(v, str) else v) for k, v in row.items()})
    return records

def parse_docx(path: str) -> List[Dict[str, Any]]:
    """Very small docx parser: extract any tables and return rows; fallback to paragraphs.
    Returns list of dicts where keys are column headers when a table is found.
    """
    try:
        import docx
    except Exception:
        raise RuntimeError("python-docx is required to parse .docx files")

    doc = docx.Document(path)
    records = []
    for table in doc.tables:
        # first row header
        headers = [c.text.strip() for c in table.rows[0].cells]
        for r in table.rows[1:]:
            vals = [c.text.strip() for c in r.cells]
            rec = {h: v for h, v in zip(headers, vals)}
            records.append(rec)
    if records:
        return records

    # fallback: split paragraphs into lines and try to parse key: value
    current = {}
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            if current:
                records.append(current)
                current = {}
            continue
        if ":" in text:
            k, v = text.split(":", 1)
            current[k.strip()] = v.strip()
    if current:
        records.append(current)
    return records
